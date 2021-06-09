from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

document.add_picture('profile_pic.jpg', width=Inches(2.0))
name = input('What is your name? ')
speak("Hello " + name + "how are you today ")
speak("what is your phone number ")
Phone_numb = input('What is your phone_numb? ')
speak("what is your email address")
Email = input('What is your email ? ')
document.add_paragraph(name + ' | ' + Phone_numb + ' | ' + Email + ' | ')

document.add_heading('About Me')
# about_me = input('Tell about yourself ')
document.add_paragraph(input('Tell about yourself '))

# Work experience
document.add_heading('Work Experience')
p = document.add_paragraph()
company = input('Enter company ')
from_date = input('From date ')
to_date = input('To date ')
p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True
experience_details = input('Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# More experiences
while True:
    more_experiences = input('Do you have more experiences? Yes or No ')

    if more_experiences.lower() == "yes":
        p = document.add_paragraph()
        company = input('Enter company ')
        from_date = input('From date ')
        to_date = input('To date ')
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True
        experience_details = input('Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break
# Skills
document.add_heading('Skills')
Skill = input('Enter skill ')
p = document.add_paragraph(Skill)
p.style = 'List Bullet'

# More Skills
while True:
    more_skills = input("Do you have other skills? Yes or No ")
    if more_skills.lower() == 'yes':
        # document.add_heading('Skills')
        Skill = input('Enter skill ')
        p = document.add_paragraph(Skill)
        p.style = 'List Bullet'
    else:
        break
print('---------------------------')
print("Bye bye................")

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Amigoscode and Brighttech Innovation"
document.save('cv.docx')
